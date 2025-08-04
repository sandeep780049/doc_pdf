from flask import Flask, render_template, request, send_file, redirect, url_for, session
import os
from werkzeug.utils import secure_filename
from PyPDF2 import PdfMerger, PdfReader
from fpdf import FPDF
from PIL import Image
from docx import Document
import uuid

app = Flask(__name__)
app.secret_key = "your_secret_key"
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    uploaded_file = request.files['file']
    if uploaded_file.filename.endswith('.docx'):
        doc = Document(uploaded_file)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for para in doc.paragraphs:
            pdf.multi_cell(0, 10, para.text)
        filename = f"{uuid.uuid4()}.pdf"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        pdf.output(output_path)
        return send_file(output_path, as_attachment=True)
    elif uploaded_file.filename.endswith(('.png', '.jpg', '.jpeg')):
        image = Image.open(uploaded_file).convert('RGB')
        filename = f"{uuid.uuid4()}.pdf"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        image.save(output_path)
        return send_file(output_path, as_attachment=True)
    elif uploaded_file.filename.endswith('.pdf'):
        filename = f"{uuid.uuid4()}.docx"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        document = Document()
        reader = PdfReader(uploaded_file)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                document.add_paragraph(text)
        document.save(output_path)
        return send_file(output_path, as_attachment=True)
    else:
        return "Unsupported file format."

@app.route('/merge', methods=['GET', 'POST'])
def merge():
    if request.method == 'POST':
        files = request.files.getlist('pdfs')
        merger = PdfMerger()
        for f in files:
            merger.append(f)
        filename = f"{uuid.uuid4()}_merged.pdf"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        merger.write(output_path)
        merger.close()
        return send_file(output_path, as_attachment=True)
    return render_template('merge.html')

@app.route('/split', methods=['GET', 'POST'])
def split():
    if request.method == 'POST':
        pdf_file = request.files['pdf']
        reader = PdfReader(pdf_file)
        for i, page in enumerate(reader.pages):
            writer = PdfMerger()
            writer.append(reader, pages=(i, i+1))
            output_filename = f"{uuid.uuid4()}_page_{i+1}.pdf"
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            writer.write(output_path)
            writer.close()
            return send_file(output_path, as_attachment=True)
    return render_template('split.html')

@app.route('/history')
def history():
    files = os.listdir(app.config['UPLOAD_FOLDER'])
    return render_template('history.html', files=files)

@app.route('/download/<filename>')
def download(filename):
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    return send_file(path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)